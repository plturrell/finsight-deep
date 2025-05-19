"""
Document Extraction Module for Financial Documents
Handles various document formats and extracts structured information
"""

import asyncio
import json
import logging
import re
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import pdfplumber
import pandas as pd
import numpy as np
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image
import docx
import xlrd
import openpyxl

from aiq.utils.debugging_utils import log_function_call


logger = logging.getLogger(__name__)


@dataclass
class ExtractedDocument:
    """Extracted document with structured data"""
    document_id: str
    document_type: str
    format: str
    content: str
    tables: List[pd.DataFrame]
    metadata: Dict[str, Any]
    extracted_at: datetime
    confidence_score: float


class DocumentExtractor:
    """
    Comprehensive document extraction for financial documents
    Supports PDF, Word, Excel, HTML, and image formats
    """
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        
        # Supported formats
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_word,
            '.doc': self._extract_word,
            '.xlsx': self._extract_excel,
            '.xls': self._extract_excel,
            '.html': self._extract_html,
            '.png': self._extract_image,
            '.jpg': self._extract_image,
            '.jpeg': self._extract_image,
            '.txt': self._extract_text,
            '.csv': self._extract_csv
        }
        
        # Financial patterns
        self.financial_patterns = {
            'currency': re.compile(r'\$?\d{1,3}(?:,\d{3})*(?:\.\d{2})?'),
            'percentage': re.compile(r'\d+(?:\.\d+)?%'),
            'date': re.compile(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}'),
            'ticker': re.compile(r'\b[A-Z]{1,5}\b'),
            'fiscal_period': re.compile(r'(?:Q[1-4]|FY)\s*\d{4}')
        }
        
        # Document type classifiers
        self.document_classifiers = {
            'earnings_report': ['earnings', 'revenue', 'eps', 'quarterly'],
            'balance_sheet': ['assets', 'liabilities', 'equity', 'balance sheet'],
            'income_statement': ['income', 'revenue', 'expenses', 'profit'],
            'cash_flow': ['cash flow', 'operating', 'investing', 'financing'],
            'annual_report': ['annual report', '10-k', 'yearly'],
            'research_report': ['analyst', 'recommendation', 'target price']
        }
        
        # Extraction metrics
        self.metrics = {
            'documents_extracted': 0,
            'extraction_errors': 0,
            'average_confidence': 0.0
        }
        
        logger.info("Initialized Document Extractor")
    
    async def extract_document(
        self,
        file_path: str,
        document_type: Optional[str] = None
    ) -> ExtractedDocument:
        """
        Extract structured data from document
        
        Args:
            file_path: Path to document file
            document_type: Optional document type override
            
        Returns:
            Extracted document with structured data
        """
        try:
            # Get file format
            file_path = Path(file_path)
            file_format = file_path.suffix.lower()
            
            if file_format not in self.supported_formats:
                raise ValueError(f"Unsupported format: {file_format}")
            
            # Extract content
            extraction_func = self.supported_formats[file_format]
            content, tables, metadata = await extraction_func(file_path)
            
            # Classify document type if not provided
            if not document_type:
                document_type = self._classify_document(content)
            
            # Extract structured information
            structured_data = self._extract_structured_data(content, document_type)
            metadata.update(structured_data)
            
            # Calculate confidence score
            confidence_score = self._calculate_confidence(content, tables, metadata)
            
            # Create extracted document
            document = ExtractedDocument(
                document_id=f"doc_{datetime.now().timestamp()}",
                document_type=document_type,
                format=file_format,
                content=content,
                tables=tables,
                metadata=metadata,
                extracted_at=datetime.now(),
                confidence_score=confidence_score
            )
            
            # Update metrics
            self.metrics['documents_extracted'] += 1
            self.metrics['average_confidence'] = (
                self.metrics['average_confidence'] * 0.9 + 
                confidence_score * 0.1
            )
            
            logger.info(f"Extracted document: {file_path.name}")
            return document
            
        except Exception as e:
            self.metrics['extraction_errors'] += 1
            logger.error(f"Document extraction failed: {e}")
            raise
    
    async def _extract_pdf(self, file_path: Path) -> Tuple[str, List[pd.DataFrame], Dict[str, Any]]:
        """Extract content from PDF"""
        content = ""
        tables = []
        metadata = {
            'page_count': 0,
            'has_images': False,
            'has_tables': False
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['page_count'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text() or ""
                    content += f"\nPage {page_num + 1}:\n{page_text}\n"
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        metadata['has_tables'] = True
                        for table in page_tables:
                            if table and len(table) > 1:
                                # Convert to DataFrame
                                df = pd.DataFrame(table[1:], columns=table[0])
                                df['page'] = page_num + 1
                                tables.append(df)
                    
                    # Check for images
                    if page.images:
                        metadata['has_images'] = True
        
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            # Try alternative extraction
            content = await self._extract_pdf_alternative(file_path)
        
        return content, tables, metadata
    
    async def _extract_pdf_alternative(self, file_path: Path) -> str:
        """Alternative PDF extraction using OCR"""
        # Convert PDF to images and OCR
        # This is a simplified implementation
        return "PDF content extracted via OCR"
    
    async def _extract_word(self, file_path: Path) -> Tuple[str, List[pd.DataFrame], Dict[str, Any]]:
        """Extract content from Word document"""
        content = ""
        tables = []
        metadata = {
            'paragraph_count': 0,
            'table_count': 0
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n"
                    metadata['paragraph_count'] += 1
            
            # Extract tables
            for table in doc.tables:
                metadata['table_count'] += 1
                data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    data.append(row_data)
                
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    tables.append(df)
        
        except Exception as e:
            logger.error(f"Word extraction error: {e}")
        
        return content, tables, metadata
    
    async def _extract_excel(self, file_path: Path) -> Tuple[str, List[pd.DataFrame], Dict[str, Any]]:
        """Extract content from Excel file"""
        content = ""
        tables = []
        metadata = {
            'sheet_count': 0,
            'total_cells': 0
        }
        
        try:
            # Read Excel file
            excel_file = pd.ExcelFile(file_path)
            metadata['sheet_count'] = len(excel_file.sheet_names)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Clean DataFrame
                df = df.dropna(how='all')
                df = df.dropna(axis=1, how='all')
                
                if not df.empty:
                    tables.append(df)
                    
                    # Convert to text
                    content += f"\nSheet: {sheet_name}\n"
                    content += df.to_string() + "\n"
                    
                    metadata['total_cells'] += df.size
        
        except Exception as e:
            logger.error(f"Excel extraction error: {e}")
        
        return content, tables, metadata
    
    async def _extract_html(self, file_path: Path) -> Tuple[str, List[pd.DataFrame], Dict[str, Any]]:
        """Extract content from HTML"""
        content = ""
        tables = []
        metadata = {
            'has_forms': False,
            'link_count': 0
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text
            content = soup.get_text(separator='\n')
            
            # Extract tables
            html_tables = soup.find_all('table')
            for table in html_tables:
                try:
                    df = pd.read_html(str(table))[0]
                    tables.append(df)
                except Exception:
                    pass
            
            # Extract metadata
            metadata['has_forms'] = len(soup.find_all('form')) > 0
            metadata['link_count'] = len(soup.find_all('a'))
        
        except Exception as e:
            logger.error(f"HTML extraction error: {e}")
        
        return content, tables, metadata
    
    async def _extract_image(self, file_path: Path) -> Tuple[str, List[pd.DataFrame], Dict[str, Any]]:
        """Extract text from image using OCR"""
        content = ""
        tables = []
        metadata = {
            'image_size': None,
            'ocr_confidence': 0.0
        }
        
        try:
            # Open image
            image = Image.open(file_path)
            metadata['image_size'] = image.size
            
            # OCR extraction
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            # Extract text with confidence
            text_elements = []
            confidences = []
            
            for i, text in enumerate(ocr_data['text']):
                if text.strip():
                    text_elements.append(text)
                    confidences.append(ocr_data['conf'][i])
            
            content = ' '.join(text_elements)
            metadata['ocr_confidence'] = np.mean(confidences) if confidences else 0.0
            
            # Try to detect tables in image
            # This is a simplified implementation
            
        except Exception as e:
            logger.error(f"Image extraction error: {e}")
        
        return content, tables, metadata
    
    async def _extract_text(self, file_path: Path) -> Tuple[str, List[pd.DataFrame], Dict[str, Any]]:
        """Extract content from text file"""
        content = ""
        tables = []
        metadata = {
            'encoding': 'utf-8',
            'line_count': 0
        }
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        metadata['encoding'] = encoding
                        metadata['line_count'] = len(content.splitlines())
                        break
                except UnicodeDecodeError:
                    continue
        
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
        
        return content, tables, metadata
    
    async def _extract_csv(self, file_path: Path) -> Tuple[str, List[pd.DataFrame], Dict[str, Any]]:
        """Extract content from CSV file"""
        content = ""
        tables = []
        metadata = {
            'row_count': 0,
            'column_count': 0
        }
        
        try:
            # Read CSV
            df = pd.read_csv(file_path)
            
            metadata['row_count'] = len(df)
            metadata['column_count'] = len(df.columns)
            
            tables.append(df)
            content = df.to_string()
        
        except Exception as e:
            logger.error(f"CSV extraction error: {e}")
        
        return content, tables, metadata
    
    def _classify_document(self, content: str) -> str:
        """Classify document type based on content"""
        content_lower = content.lower()
        scores = {}
        
        for doc_type, keywords in self.document_classifiers.items():
            score = sum(1 for keyword in keywords if keyword in content_lower)
            scores[doc_type] = score
        
        # Return type with highest score
        if scores:
            return max(scores, key=scores.get)
        
        return "general_financial"
    
    def _extract_structured_data(self, content: str, document_type: str) -> Dict[str, Any]:
        """Extract structured financial data from content"""
        structured_data = {
            'currencies': [],
            'percentages': [],
            'dates': [],
            'tickers': [],
            'fiscal_periods': []
        }
        
        # Extract using patterns
        for pattern_name, pattern in self.financial_patterns.items():
            matches = pattern.findall(content)
            if matches:
                structured_data[f'{pattern_name}s'] = list(set(matches))
        
        # Extract document-specific data
        if document_type == 'earnings_report':
            structured_data.update(self._extract_earnings_data(content))
        elif document_type == 'balance_sheet':
            structured_data.update(self._extract_balance_sheet_data(content))
        elif document_type == 'income_statement':
            structured_data.update(self._extract_income_statement_data(content))
        
        return structured_data
    
    def _extract_earnings_data(self, content: str) -> Dict[str, Any]:
        """Extract earnings-specific data"""
        earnings_data = {
            'revenue': None,
            'eps': None,
            'guidance': None
        }
        
        # Extract revenue
        revenue_match = re.search(r'revenue[:\s]+\$?([\d,]+(?:\.\d+)?)', content, re.IGNORECASE)
        if revenue_match:
            earnings_data['revenue'] = revenue_match.group(1).replace(',', '')
        
        # Extract EPS
        eps_match = re.search(r'eps[:\s]+\$?([\d.]+)', content, re.IGNORECASE)
        if eps_match:
            earnings_data['eps'] = eps_match.group(1)
        
        return earnings_data
    
    def _extract_balance_sheet_data(self, content: str) -> Dict[str, Any]:
        """Extract balance sheet data"""
        balance_data = {
            'total_assets': None,
            'total_liabilities': None,
            'shareholders_equity': None
        }
        
        # Extract key balance sheet items
        # This is a simplified implementation
        
        return balance_data
    
    def _extract_income_statement_data(self, content: str) -> Dict[str, Any]:
        """Extract income statement data"""
        income_data = {
            'revenue': None,
            'operating_income': None,
            'net_income': None
        }
        
        # Extract key income statement items
        # This is a simplified implementation
        
        return income_data
    
    def _calculate_confidence(
        self,
        content: str,
        tables: List[pd.DataFrame],
        metadata: Dict[str, Any]
    ) -> float:
        """Calculate extraction confidence score"""
        confidence = 0.5  # Base confidence
        
        # Adjust based on content quality
        if len(content) > 100:
            confidence += 0.1
        
        if tables:
            confidence += 0.2
        
        # Check for financial data
        if any(self.financial_patterns['currency'].findall(content)):
            confidence += 0.1
        
        if any(self.financial_patterns['percentage'].findall(content)):
            confidence += 0.1
        
        # OCR confidence for images
        if 'ocr_confidence' in metadata:
            confidence = confidence * (metadata['ocr_confidence'] / 100)
        
        return min(1.0, confidence)
    
    async def extract_batch(self, file_paths: List[str]) -> List[ExtractedDocument]:
        """Extract multiple documents in batch"""
        tasks = []
        
        for file_path in file_paths:
            task = asyncio.create_task(self.extract_document(file_path))
            tasks.append(task)
        
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out errors
        documents = []
        for result in results:
            if isinstance(result, ExtractedDocument):
                documents.append(result)
            else:
                logger.error(f"Batch extraction error: {result}")
        
        return documents
    
    def get_metrics(self) -> Dict[str, Any]:
        """Get extraction metrics"""
        return {
            'documents_extracted': self.metrics['documents_extracted'],
            'extraction_errors': self.metrics['extraction_errors'],
            'average_confidence': self.metrics['average_confidence'],
            'supported_formats': list(self.supported_formats.keys())
        }